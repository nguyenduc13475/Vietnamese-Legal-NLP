import io
import os
import re
import tempfile
import time
import unicodedata

import aspose.words as aw
import docx
from google import genai
from google.genai import types
from pypdf import PdfReader

from src.utils.prompts import DOCUMENT_CLEANING_PROMPT


def clean_document_with_gemini(filepath: str, api_key: str) -> str:
    """
    Converts non-PDF documents to PDF, uploads to Gemini for cleaning,
    and returns the structured text.
    """
    client = genai.Client(api_key=api_key)
    temp_pdf_path = None
    uploaded_file = None
    upload_target = filepath

    try:
        # Convert to PDF if the file is not already a PDF
        if not filepath.lower().endswith(".pdf"):
            doc = aw.Document(filepath)
            fd, temp_pdf_path = tempfile.mkstemp(suffix=".pdf")
            os.close(fd)
            doc.save(temp_pdf_path)
            upload_target = temp_pdf_path

        # Upload to Gemini Cloud
        uploaded_file = client.files.upload(file=upload_target)

        # Generate cleaned content with multi-model fallback and retry
        models_to_try = [
            "gemini-3.1-flash-lite-preview",
            "gemini-2.5-flash",
            "gemini-3-flash-preview",
            "gemini-2.5-flash-lite-preview-09-2025",
            "gemma-3-27b-it",
        ]
        max_full_rolls = 5
        last_error = None
        response = None

        for roll in range(max_full_rolls):
            for model_name in models_to_try:
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents=[DOCUMENT_CLEANING_PROMPT, uploaded_file],
                        config=types.GenerateContentConfig(temperature=0.4),
                    )
                    break  # Success, break inner loop
                except Exception as e:
                    last_error = e
                    print(
                        f"[Document Parser Warning] Model {model_name} failed. Retrying in 5s... Error: {str(e)}"
                    )
                    time.sleep(5)

            if response:
                break  # Success, break outer loop

            if roll < max_full_rolls - 1:
                print(
                    f"[Document Parser Warning] All models failed in roll {roll + 1}/{max_full_rolls}. Waiting 60s before next roll..."
                )
                time.sleep(60)

        if not response:
            raise Exception(
                f"Gemini API completely failed to parse document after {max_full_rolls} full rolls. Last error: {str(last_error)}"
            )

        return response.text.strip()

    finally:
        # Cleanup Gemini Cloud file
        if uploaded_file:
            try:
                client.files.delete(name=uploaded_file.name)
            except Exception:
                pass

        # Cleanup local temporary PDF
        if temp_pdf_path and os.path.exists(temp_pdf_path):
            try:
                os.remove(temp_pdf_path)
            except Exception:
                pass


def extract_text_from_txt(content: bytes) -> str:
    return content.decode("utf-8", errors="ignore")


def extract_text_from_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(content))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        return f"[Error] Lỗi khi đọc file PDF: {str(e)}"


def extract_text_from_docx(content: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(content))
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        return f"[Error] Error reading DOCX file: {str(e)}"


def clean_contract_text(text: str) -> str:
    """
    Advanced text cleaning specifically tuned for Vietnamese Legal Contracts.
    """
    # 1. Unicode Normalization (NFC)
    text = unicodedata.normalize("NFC", text)

    # 2. Remove hidden characters
    text = re.sub(r"[\u200B-\u200D\uFEFF\u00AD]", "", text)

    # 3. Remove Table of Contents (TOC)
    text = re.sub(r"(?m)^.*\.{4,}\s*\d+\s*$", "", text)

    # 4. Remove page numbering
    text = re.sub(
        r"(?i)^\s*(?:page|trang)\s*\d+\s*(?:of|/)?\s*\d*\s*$",
        "",
        text,
        flags=re.MULTILINE,
    )
    text = re.sub(r"^\s*-\s*\d+\s*-\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"(?m)^\s*[-_=]{3,}\s*$", "", text)

    # 5. Whitespace Normalization
    text = re.sub(r"[ \t]+", " ", text)

    lines = text.split("\n")
    cleaned_lines = []
    bullet_pattern = re.compile(
        r"^([a-zđ][\)\.]|\d+(?:\.\d+)*[\.\):]?|[IVXLCDM]+[\.\)]|[-+]|Điều\s+\d+|Khoản\s+\d+|Mục\s+\d+|Phần\s+\d+)",
        re.IGNORECASE,
    )

    for line in lines:
        line = line.strip()
        if not line:
            cleaned_lines.append("")
            continue

        if cleaned_lines and cleaned_lines[-1] != "":
            prev_line = cleaned_lines[-1]
            ends_with_punctuation = re.search(r"[.:;!?]$", prev_line)
            is_all_caps = prev_line.isupper()
            starts_with_bullet = bullet_pattern.match(line)
            starts_with_lower_or_comma = re.match(
                r"^[,\s]*[a-zđáàảãạâấầẩẫậăắằẳẵặéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵ]",
                line,
            )
            ends_with_conjunction = re.search(
                r"(?i)(?:và|hoặc|của|cho|bởi|như|là|thì|mà|để|với|tại|từ|về|thuộc)\s*$",
                prev_line,
            )

            if not ends_with_punctuation and not is_all_caps and not starts_with_bullet:
                if starts_with_lower_or_comma or ends_with_conjunction:
                    cleaned_lines[-1] = prev_line + " " + line
                    continue

        cleaned_lines.append(line)

    final_text = "\n".join(cleaned_lines)
    final_text = re.sub(r"\n{3,}", "\n\n", final_text)
    return final_text.strip()


def parse_and_clean_document(file_content: bytes, filename: str) -> str:
    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        text = extract_text_from_pdf(file_content)
    elif ext in ["doc", "docx"]:
        text = extract_text_from_docx(file_content)
    else:
        text = extract_text_from_txt(file_content)

    return clean_contract_text(text)
